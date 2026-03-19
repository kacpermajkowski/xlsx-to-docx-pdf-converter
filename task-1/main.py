import signal
import tempfile
from datetime import datetime
from pathlib import Path
import threading
import time
import webbrowser
import multiprocessing

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

from docx2pdf import convert

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from flask import Flask, render_template, request, send_file
import pandas
import os.path

import zipfile
import io

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS  # PyInstaller temp folder
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

app = Flask(
    __name__,
    template_folder=resource_path("templates"),
    static_folder=resource_path("static")
)


@app.route('/', methods=['GET'])
def index():
    return render_template('index.html', output="test")


# Parameter list:
# - file: the uploaded Excel file
# Form params:
# - mode: the output format (pdf, docx, both)
# - paragraph_spacing: the spacing between paragraphs in the output document (in points)
# - line_spacing: the spacing between lines in the output document (e.g., 1.5 for 1.5x line spacing)
# - title: the title of the output document (optional)
# - alignment: the text alignment in the output document (left, center, right, justify)
# - page_numbers: whether to include page numbers in the output document (yes or null)
# - font_size: the font size in the output document (in points)
#
#  request.form['name']
#  request.form.get('name', 'default_value')  # Use this to avoid KeyError if the key is missing

@app.route('/upload-file', methods=['POST'])
def upload_file():
    file = request.files['file']
    if not file or file.filename == '':
        return "No file uploaded", 400
    if not file.filename.endswith('.xlsx'):
        return "Invalid file type. Supported types: .xlsx", 400

    df = pandas.read_excel(file)
    df = df.replace(r'\r\n|\r|\n', ' ', regex=True)
    base = Path(file.filename).stem
    format = request.form['format']
    mode = request.form.get('mode', 'document')
    if format == 'pdf':
        pdf_buffer = None
        if mode == 'document':
            pdf_buffer = convert_to_pdf(base, df)
        if mode == 'table':
            pdf_buffer = df_to_pdf_table(base, df)
        return send_file(pdf_buffer,
                         mimetype='application/pdf',
                         as_attachment=True,
                         download_name=base + '.pdf')
    elif format == 'docx':
        docx_buffer = None
        if mode == 'document':
            docx_buffer = convert_to_docx(base, df)
        if mode == 'table':
            docx_buffer = df_to_docx_table(base, df)
        return send_file(docx_buffer,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True,
                         download_name=base + '.docx')
    elif format == 'both':
        pdf_buffer = None
        docx_buffer = None
        if mode == 'document':
            docx_buffer = convert_to_docx(base, df)
            pdf_buffer = convert_to_pdf(base, df)
        if mode == 'table':
            docx_buffer = df_to_docx_table(base, df)
            pdf_buffer = df_to_pdf_table(base, df)

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(base + '.pdf', pdf_buffer.getvalue())
            zf.writestr(base + '.docx', docx_buffer.getvalue())

        zip_buffer.seek(0)

        return send_file(zip_buffer,
                         mimetype='application/zip',
                         as_attachment=True,
                         download_name=base + '.zip')
    else:
        return "Invalid mode. Supported modes: pdf, docx, both", 400

def convert_to_pdf(base, df):
    align_map = {
        "left": TA_LEFT,
        "center": TA_CENTER,
        "right": TA_RIGHT,
        "justify": TA_JUSTIFY,
    }
        
    pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Bold', 'C:/Windows/Fonts/arialbd.ttf'))
    pdfmetrics.registerFontFamily('Arial', normal='Arial', bold='Arial-Bold')

    buffer = io.BytesIO()

    paragraph_spacing = int(request.form['paragraph_spacing'])
    line_spacing = float(request.form['line_spacing'])
    title = request.form.get('title', '')
    alignment = align_map.get(request.form['alignment'].lower(), TA_LEFT)
    page_numbers = request.form.get('page_numbers', 'no') == 'yes'
    font_size = int(request.form['font_size'])

    styles = getSampleStyleSheet()

    normal_style = ParagraphStyle(
        'Custom',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=font_size,
        alignment=alignment,
        spaceBefore=paragraph_spacing,
        spaceAfter=paragraph_spacing,
        leading=line_spacing * font_size
    )
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=font_size,
        leading=font_size * line_spacing,
    )
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontName='Arial-Bold',
        fontSize=24,
        alignment=TA_CENTER,
    )

    data_style = ParagraphStyle(
        'CustomData',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=font_size,
        alignment=TA_CENTER,
    )

    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []

    if title:
        story.append(Spacer(1, 80))
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph('Data wygenerowania: ' + datetime.now().strftime("%Y-%m-%d %H:%M:%S"), data_style))
        story.append(PageBreak())

    for i, row in df.iterrows():
        if page_numbers:
            hp = f"{base} | Row {i + 1} of {len(df)}"
            story.append(Paragraph(hp, header_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black, spaceAfter=6))

        for col in df.columns:
            line = f"<b>{col}:</b> {row[col]}"
            story.append(Paragraph(line, normal_style))

        if i < len(df) - 1:
            story.append(PageBreak())

        story.append(PageBreak())

    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_to_docx(base, df):
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    
    document = Document()

    title = request.form['title']

    alignment = align_map.get(request.form['alignment'].lower(), WD_ALIGN_PARAGRAPH.LEFT)
    spacing_pt = Pt(int(request.form['paragraph_spacing']))
    line_spacing = float(request.form['line_spacing'])

    if title:
        title_paragraph = document.add_paragraph()
        title_paragraph.paragraph_format.space_before = Pt(120)
        run = title_paragraph.add_run(title)
        run.bold = True
        run.font.size = Pt(24)
        run.font.name = 'Arial'
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('Data wygenerowania: ' + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_break(WD_BREAK.PAGE)

    for i, row in df.iterrows():
        if request.form.get('page_numbers', 'no') == 'yes':
            hp = document.add_paragraph(base + ' | Row ' + str(i + 1) + ' of ' + str(len(df)))
            insert_hr_docx(hp)
        last_p = None
        for col in df.columns:
            p = document.add_paragraph()

            p.alignment = alignment
            p.paragraph_format.space_below = spacing_pt
            p.paragraph_format.space_before = spacing_pt
            p.paragraph_format.line_spacing = line_spacing

            r = p.add_run(col + ': ')
            r.bold = True
            r.font.size = Pt(int(request.form['font_size']))
            r.font.name = 'Arial'

            r = p.add_run(str(row[col]))
            r.font.size = Pt(int(request.form['font_size']))
            r.font.name = 'Arial'

            last_p = p
        last_p.add_run().add_break(WD_BREAK.PAGE)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def df_to_docx_table(base, df):
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
        
    line_spacing = float(request.form.get('line_spacing', 1.0))
    title        = request.form.get('title', None)
    alignment    = request.form.get('alignment', 'left')
    font_size    = float(request.form.get('font_size', 11))

    align = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    def set_landscape(section):
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
            setattr(section, attr, Cm(1))

    def style_cell(cell):
        for para in cell.paragraphs:
            para.alignment = align
            para.paragraph_format.line_spacing = line_spacing
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(font_size)

    doc = Document()

    if title:
        set_landscape(doc.sections[0])

        doc.add_paragraph()

        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(font_size + 6)
            run.bold = True

        date_para = doc.add_paragraph('Data wygenerowania: ' + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in date_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(font_size)

        doc.add_page_break()

    section = doc.add_section() if title else doc.sections[0]
    set_landscape(section)

    table = doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')

    for cell, col in zip(table.rows[0].cells, df.columns):
        cell.text = str(col)
        style_cell(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for cell, val in zip(cells, row):
            cell.text = '' if val is None else str(val)
            style_cell(cell)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def df_to_pdf_table(base, df):
    docx_buffer = df_to_docx_table(base, df)

    docx_path = f"{base}.docx"
    pdf_path = f"{base}.pdf"

    try:
        with open(docx_path, 'wb') as f:
            f.write(docx_buffer.read())

        convert(docx_path, pdf_path)

        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)
        if os.path.exists(pdf_path):
            os.remove(pdf_path)

    return io.BytesIO(pdf_bytes)

# Source: https://github.com/python-openxml/python-docx/issues/105#issuecomment-442786431
def insert_hr_docx(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

@app.route('/shutdown', methods=['POST', 'GET'])
def shutdown():
    kill_app()

def kill_app():
    os.kill(multiprocessing.current_process().pid, signal.SIGTERM) 
    
def reset_timer():
    global _life_timer
    if _life_timer is not None:
        _life_timer.cancel()
    
    _life_timer = threading.Timer(60.0, kill_app)
    _life_timer.daemon = True
    _life_timer.start()
    
@app.route('/keepalive', methods=['POST', 'GET'])
def keepalive():
    reset_timer()
    return 'Timer reset', 200

def open_site():
    webbrowser.open("http://127.0.0.1:5000")

def run_flask():
    app.run(port=5000)

_life_timer = None

if __name__ == '__main__':
    threading.Thread(target=open_site).start()
    threading.Thread(target=run_flask).start()
    reset_timer()
    
    


